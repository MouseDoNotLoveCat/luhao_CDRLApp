#!/usr/bin/env python3
"""
分析 Word 文件示范的脚本
用于了解监督通知书的具体格式和结构
"""

import os
from pathlib import Path

# 尝试导入 python-docx，如果没有则提示安装
try:
    from docx import Document
    from docx.table import Table
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    exit(1)


def analyze_word_file(file_path):
    """分析 Word 文件的结构"""
    print(f"\n{'='*80}")
    print(f"分析文件: {file_path}")
    print(f"{'='*80}\n")
    
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"❌ 无法打开文件: {e}")
        return
    
    # 1. 分析段落
    print("📄 段落内容:")
    print("-" * 80)
    for i, para in enumerate(doc.paragraphs[:20]):  # 只显示前20个段落
        text = para.text.strip()
        if text:
            print(f"  [{i}] {text[:100]}")
    
    if len(doc.paragraphs) > 20:
        print(f"  ... 还有 {len(doc.paragraphs) - 20} 个段落")
    
    # 2. 分析表格
    print(f"\n📊 表格数量: {len(doc.tables)}")
    print("-" * 80)
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n  表格 {table_idx + 1}:")
        print(f"    行数: {len(table.rows)}, 列数: {len(table.columns)}")
        
        # 显示表格前几行
        for row_idx, row in enumerate(table.rows[:5]):
            cells_text = [cell.text.strip()[:20] for cell in row.cells]
            print(f"    行 {row_idx}: {cells_text}")
        
        if len(table.rows) > 5:
            print(f"    ... 还有 {len(table.rows) - 5} 行")
    
    # 3. 分析图片
    print(f"\n🖼️  图片数量: {count_images(doc)}")
    print("-" * 80)
    
    # 4. 分析文本特征
    print(f"\n🔍 文本特征分析:")
    print("-" * 80)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # 查找关键字段
    keywords = [
        "项目名称", "标段", "工点", "建设单位", "施工单位", 
        "监理单位", "检查时间", "检查人员", "问题", "整改"
    ]
    
    for keyword in keywords:
        if keyword in full_text:
            print(f"  ✅ 找到关键字: {keyword}")
        else:
            print(f"  ❌ 未找到关键字: {keyword}")
    
    # 5. 统计信息
    print(f"\n📈 统计信息:")
    print("-" * 80)
    print(f"  总段落数: {len(doc.paragraphs)}")
    print(f"  总表格数: {len(doc.tables)}")
    print(f"  总文本长度: {len(full_text)} 字符")
    print(f"  总行数: {len(full_text.split(chr(10)))}")


def count_images(doc):
    """计算文档中的图片数量"""
    count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            count += 1
    return count


def main():
    """主函数"""
    samples_dir = Path("Samples")
    
    if not samples_dir.exists():
        print(f"❌ 找不到 Samples 目录")
        return
    
    # 查找所有 .docx 文件
    docx_files = list(samples_dir.glob("*.docx"))
    
    if not docx_files:
        print(f"❌ 在 Samples 目录中找不到 .docx 文件")
        return
    
    print(f"\n🔍 找到 {len(docx_files)} 个 Word 文件")
    
    for file_path in docx_files:
        analyze_word_file(str(file_path))
    
    print(f"\n{'='*80}")
    print("✅ 分析完成")
    print(f"{'='*80}\n")


if __name__ == "__main__":
    main()

