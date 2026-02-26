# -*- coding: utf-8 -*-
"""
Word 文档解析模块
用于解析监督通知书 Word 文档
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from docx import Document
from datetime import datetime


class WordDocumentParser:
    """Word 文档解析器"""

    def __init__(self, file_path: str):
        """
        初始化解析器

        Args:
            file_path: Word 文件路径
        """
        self.file_path = Path(file_path)
        self.doc = None
        self.paragraphs = []
        self.current_section = None
        self.document_structure = None  # 文档结构类型：'two_level' 或 'three_level'

    def parse(self) -> Dict:
        """
        解析 Word 文档

        Returns:
            解析结果字典
        """
        try:
            self.doc = Document(str(self.file_path))
            self._extract_paragraphs()

            # 自动检测文档结构（二级 vs 三级）
            self.document_structure = self._detect_document_structure()

            result = {
                'file_name': self.file_path.name,
                'notice_number': self._extract_notice_number(),
                'check_date': self._extract_check_date(),
                'builder_unit': self._extract_builder_unit(),
                'inspection_unit': self._extract_inspection_unit_from_first_para(),
                'inspection_personnel': self._extract_inspection_personnel_from_first_para(),
                'inspection_basis': self._extract_inspection_basis(),
                'project_name': self._extract_project_name_from_first_para() or self._extract_project_name(),
                'check_unit': self._extract_check_unit() or '未知单位',
                'check_personnel': self._extract_check_personnel(),
                'project_name_old': self._extract_project_name(),
                'rectification_notices': self._extract_rectification_notices(),
                'other_issues': self._extract_other_issues(),
                'total_issues': 0,
                'declared_issues_count': None,
                'warnings': [],
                'status': 'success',
                'document_structure': self.document_structure  # 添加结构信息到结果中
            }

            # 统一为每条问题补全施工/监理单位（基于标段编号映射与上下文回填）
            self._postprocess_fill_units(result)

            result['total_issues'] = len(result['rectification_notices']) + len(result['other_issues'])

            # 验证问题总数
            declared_count_info = self._extract_total_issues_count()
            if declared_count_info:
                result['declared_issues_count'] = declared_count_info
                declared_count = declared_count_info.get('total')
                actual_count = result['total_issues']

                if declared_count != actual_count:
                    result['warnings'].append(
                        f'问题总数不匹配：文档声明 {declared_count} 个，实际识别 {actual_count} 个'
                    )

            return result

        except Exception as e:
            return {
                'file_name': self.file_path.name,
                'status': 'error',
                'error': str(e)
            }

    def _extract_paragraphs(self):
        """提取所有段落"""
        # 保存段落对象和文本，以便后续可以访问段落的格式属性（如Word自动编号）
        self.paragraph_objects = [p for p in self.doc.paragraphs if p.text.strip()]
        self.paragraphs = [p.text.strip() for p in self.paragraph_objects]

    def _detect_document_structure(self) -> str:
        """
        自动检测文档的层级结构

        返回值：
        - 'two_level': 二级结构（标段+问题），标段行中包含具体工点名称
        - 'three_level': 三级结构（标段+工点+问题），标段行中不包含具体工点名称

        检测逻辑：
        1. 找到第一个标段行（以（一）、（二）等开头，包含"施工"和"监理"）
        2. 检查该标段行中是否包含具体工点名称（如桥梁名、隧道名、站场名等）
        3. 根据是否包含工点名称来判断结构类型
        4. 若文档使用数字序号（1. 2. 3.）而非中文序号，同样检测其标段行
        """
        # 工点名称的特征词汇（用于识别具体工点名称）
        site_name_keywords = [
            '桥', '隧道', '站', '路基', '基坑', '挡墙', '边坡', '排水', '防护',
            '梁', '墩', '拱', '涵', '通道', '通路', '斜井', '竖井', '出口', '入口',
            '接触网', '信号', '通信', '电力', '给水', '污水', '雨水', '燃气',
            '大桥', '特大桥', '中桥', '小桥', '高架', '地下', '地面',
            # 补充：实际文档中出现的工点类型
            '普速场', '信号楼', '内业资料', '场坪', '站场', '货场', '机务段',
            '动车所', '变电所', '综合楼', '办公楼', '宿舍楼', '食堂',
        ]

        # 查找第一个标段行（中文序号格式：（一）...施工...监理）
        for para in self.paragraphs:
            if re.match(r'^（[一二三四五六七八九十]）', para) and '施工' in para and '监理' in para:
                # 检查标段行中是否包含工点名称关键词，或包含"标"后面跟着工点名称
                has_site_name = any(keyword in para for keyword in site_name_keywords)
                # 额外检测：标段编号后面是否还有内容（即工点名称）
                # 格式：...的XXXX标[工点名称]（检查时间...）
                if not has_site_name:
                    has_site_name = bool(re.search(r'[A-Z]{2,}[A-Z0-9]*(?:-?\d+)?标.+（检查', para))

                if has_site_name:
                    return 'two_level'
                else:
                    return 'three_level'

        # 查找数字序号格式的标段行（1. 2. 3. 格式，包含施工和监理信息）
        for para in self.paragraphs:
            if re.match(r'^\d+[\.．]', para) and '施工' in para and '监理' in para and '标' in para:
                # 数字序号格式的标段行，工点名称直接包含在行内 → 二级结构
                return 'two_level'

        # 如果没有找到标段行，默认返回三级结构
        return 'three_level'

    def _extract_notice_number(self) -> Optional[str]:
        """
        提取通知书编号

        格式示例：
        - 南宁站〔2025〕（通知）玉岑08号
        - 南宁站[2025]（通知）柳梧10号
        - 宁建监2025-11
        """
        # 查找编号模式
        patterns = [
            r'南宁站[〔\[]2025[〕\]]\（通知\）\S+\d+号',
            r'宁建监\d{4}-\d+',
            r'编号[：:]\s*(\S+)',
        ]

        for para in self.paragraphs[:20]:  # 只查看前 20 段
            for pattern in patterns:
                match = re.search(pattern, para)
                if match:
                    return match.group(0)

        return None

    def _extract_check_date(self) -> Optional[str]:
        """
        提取检查日期

        格式示例：2025-08-07, 2025年8月7日
        """
        date_patterns = [
            r'(\d{4})[年-](\d{1,2})[月-](\d{1,2})[日]?',
        ]

        for para in self.paragraphs[:30]:
            for pattern in date_patterns:
                match = re.search(pattern, para)
                if match:
                    year, month, day = match.groups()
                    return f"{year}-{month.zfill(2)}-{day.zfill(2)}"

        return None

    def _extract_check_unit(self) -> Optional[str]:
        """
        提取检查单位

        通常在"检查单位"或"监督单位"之后
        """
        for i, para in enumerate(self.paragraphs[:20]):
            if '检查单位' in para or '监督单位' in para:
                # 尝试从当前或下一段提取
                if i + 1 < len(self.paragraphs):
                    return self.paragraphs[i + 1]
                # 或从当前段提取冒号后的内容
                if '：' in para or ':' in para:
                    parts = re.split('[：:]', para)
                    if len(parts) > 1:
                        return parts[1].strip()

        return None

    def _extract_check_personnel(self) -> Optional[str]:
        """
        提取检查人员
        """
        for i, para in enumerate(self.paragraphs[:20]):
            if '检查人员' in para or '监督人员' in para:
                if i + 1 < len(self.paragraphs):
                    return self.paragraphs[i + 1]
                if '：' in para or ':' in para:
                    parts = re.split('[：:]', para)
                    if len(parts) > 1:
                        return parts[1].strip()

        return None

    def _extract_project_name(self) -> Optional[str]:
        """
        提取项目名称（fallback 方法）

        通常在文档开头叙述段落或"总体情况"段落中
        示例：合湛铁路、黄百铁路、钦州东至钦州港增建二线工程、钦防增建二线等

        识别规则（优先级从高到低）：
        1. 从叙述段落中匹配"对XXX工程/项目"格式（支持不含"铁路"的项目名）
           适用于：钦港二线（"对钦州东至钦州港增建二线工程项目实施监督检查"）
                   钦防二线（"对钦防增建二线项目开展了监督检查"）
        2. 优先匹配"新建xxx铁路"格式（总体情况段落中常见）
        3. 跳过含建设单位特征词的段落（"指挥部"、"有限责任公司"、"集团有限公司"等）
        4. 最后才从普通段落中提取"xxx铁路"
        """
        # 建设单位特征词，含这些词的段落不作为项目名称来源
        _BUILDER_KEYWORDS = ('指挥部', '有限责任公司', '集团有限公司', '铁路局', '建设指挥')
        # 标段编号模式（用于排除误匹配）
        _SECTION_PAT = r'[A-Z]{2,}[A-Z0-9]*(?:-?\d+)?标'

        # 第一轮：从叙述段落中匹配"对XXX工程/项目"格式
        # 匹配规则：
        #   - "对" 后跟项目名称（至少4个字，不含句号/逗号/空格）
        #   - 项目名称以"工程"或"项目"结尾
        #   - 后接动词（实施/开展/进行/开始）、数字（如"2个施工标段"）或标段编号
        # 注意：不跳过含建设单位词的段落（[03]段可能同时含"集团有限公司"和项目名）
        #       只跳过**以建设单位词结尾**的短段落（如"集团公司沿海铁路工程建设指挥部："）
        for para in self.paragraphs[:10]:
            # 跳过纯建设单位行（短段落且以指挥部/公司结尾）
            if len(para) < 30 and any(para.rstrip('：:').endswith(kw) for kw in _BUILDER_KEYWORDS):
                continue
            match = re.search(
                r'对([^，。；\s对]{4,30}(?:工程|项目))'
                r'(?=\s*(?:实施|开展|进行|开始|\d|的\s*' + _SECTION_PAT + r'))',
                para
            )
            if match:
                return match.group(1).strip()

        # 第二轮：优先匹配"新建xxx铁路"格式
        for para in self.paragraphs[:30]:
            if '铁路' in para:
                match = re.search(r'新建([^，。；\s]+铁路)', para)
                if match:
                    return match.group(1)

        # 第三轮：跳过建设单位段落，匹配普通"xxx铁路"
        for para in self.paragraphs[:30]:
            if '铁路' in para:
                if any(kw in para for kw in _BUILDER_KEYWORDS):
                    continue
                match = re.search(r'(\S+铁路)', para)
                if match:
                    return match.group(1)

        return None

    def _extract_builder_unit(self) -> Optional[str]:
        """
        提取建设单位

        位置：通知书编号的下一行
        识别规则：
        1. 查找编号行（包含"编号"、"〔"或"["）
        2. 在编号后的 1-3 行查找建设单位
        3. 建设单位通常以"指挥部"或"公司"结尾
        4. 如果包含冒号，提取冒号前的部分

        示例：
        - 柳州铁路工程建设指挥部
        - 广西铁路投资集团有限公司
        - 云桂铁路广西有限责任公司：
        """
        # 通常在编号后的 1-3 行
        for i, para in enumerate(self.paragraphs[:10]):
            # 查找编号
            if '编号' in para or '〔' in para or '[' in para:
                # 在编号后的 1-3 行查找建设单位
                for j in range(i + 1, min(i + 4, len(self.paragraphs))):
                    next_para = self.paragraphs[j]
                    # 查找"指挥部"或"公司"
                    if '指挥部' in next_para or '公司' in next_para:
                        # 提取单位名称（去掉冒号后的内容）
                        if '：' in next_para or ':' in next_para:
                            parts = re.split('[：:]', next_para)
                            return parts[0].strip()
                        else:
                            return next_para.strip()

        return None

    def _extract_inspection_unit_from_first_para(self) -> Optional[str]:
        """
        从第一段话中提取检查单位

        句子结构：`南宁监督站****对*****铁路****标****、****、****、****，****标*******、*******等工点`

        识别规则：
        1. 检查单位通常是"南宁监督站"或类似的监督站名称
        2. 位于第一段话中

        示例：
        南宁监督站蒋德义、卢浩对柳梧铁路...
        """
        # 查找第一段话（通常在前 10 段）
        for para in self.paragraphs[:10]:
            # 查找"监督站"关键词
            if '监督站' in para:
                # 提取"xxx监督站"（只提取到"监督站"为止）
                match = re.search(r'([^，。；\s]+监督站)', para)
                if match:
                    return match.group(1)

        return None

    def _extract_inspection_personnel_from_first_para(self) -> Optional[str]:
        """
        从第一段话中提取检查人员

        句子结构：
        - 情况1：`南宁监督站【检查人员】根据...对*****铁路...`
        - 情况2：`按照...南宁监督站【检查人员】对*****铁路...`

        识别规则：
        1. 检查人员位于"监督站"之后
        2. 通常是人名列表，用、分隔
        3. 在"根据"/"按照"或"对"处截断，只提取人名部分
        4. 优先在"根据"/"按照"处截断

        示例：
        - 南宁监督站【卢浩、陈胜及建设部第四检查组胡云龙】根据《...》...对黄百铁路...
        - 按照...南宁监督站【唐小林、罗斌、蒋德义】对柳梧铁路...
        """
        # 查找第一段话
        for para in self.paragraphs[:10]:
            # 查找"监督站"
            if '监督站' in para:
                # 查找"监督站"的位置
                station_pos = para.find('监督站')
                after_station = para[station_pos + 3:]

                # 查找"根据"或"按照"的位置
                basis_keywords = ['根据', '按照']
                basis_pos = -1
                for keyword in basis_keywords:
                    pos = after_station.find(keyword)
                    if pos != -1 and (basis_pos == -1 or pos < basis_pos):
                        basis_pos = pos

                # 查找"对"的位置
                dui_pos = after_station.find('对')

                # 确定截断位置：优先使用"根据"/"按照"，否则使用"对"
                end_pos = -1
                if basis_pos != -1:
                    end_pos = basis_pos
                elif dui_pos != -1:
                    end_pos = dui_pos

                if end_pos != -1:
                    # 提取"监督站"和截断位置之间的文字
                    personnel = after_station[:end_pos].strip()
                    # 去掉可能的空格和特殊字符
                    personnel = re.sub(r'[\s\u3000]+', '', personnel)
                    if personnel:
                        return personnel

        return None

    def _extract_inspection_basis(self) -> Optional[str]:
        """
        从第一段话中提取检查依据

        句子结构：`...对*****铁路...根据《...》、《...》...`

        识别规则：
        1. 检查依据位于"根据"或"按照"之后
        2. 通常包含多个文件名称和文号
        3. 以"等文件的要求"、"等文件"、"等规定"、"等要求"等词语结尾
        4. 或者以"，"后面跟着新的语义段落结尾

        示例：
        根据《国铁集团关于开展在建铁路桥梁施工安全隐患排查整治的紧急通知》(铁建设电〔2025〕44号)、
        《国家铁路局综合司关于开展铁路桥梁工程质量安全问题隐患排查治理的通知》（国铁综工程监函[2025]351号）、
        《中国铁路南宁局集团有限公司关于集中开展铁路建设项目安全隐患排查整治工作的通知》（宁铁建函〔2025〕243号）
        """
        # 查找第一段话
        for para in self.paragraphs[:10]:
            # 查找"根据"或"按照"
            if '根据' in para or '按照' in para:
                # 查找"根据"或"按照"的位置
                basis_start = -1
                if '根据' in para:
                    basis_start = para.find('根据')
                if '按照' in para and (basis_start == -1 or para.find('按照') < basis_start):
                    basis_start = para.find('按照')

                if basis_start != -1:
                    # 从"根据"或"按照"之后开始提取
                    basis_text = para[basis_start:]

                    # 查找结束标志
                    # 优先查找"等文件的要求"、"等文件"、"等规定"、"等要求"
                    end_patterns = [
                        r'等文件的要求',
                        r'等文件',
                        r'等规定',
                        r'等要求',
                        r'等通知',
                        r'等文件精神',
                    ]

                    for pattern in end_patterns:
                        match = re.search(pattern, basis_text)
                        if match:
                            # 提取到匹配位置的末尾
                            basis_text = basis_text[:match.end()]
                            break
                    else:
                        # 如果没有找到明确的结束标志，查找下一个"，"后面是否有新的语义段落
                        # 通常新段落会以"对"、"在"、"为"等词开头
                        comma_matches = list(re.finditer(r'，', basis_text))
                        if comma_matches:
                            # 检查最后一个逗号之后是否有新的语义段落
                            last_comma_pos = comma_matches[-1].end()
                            after_comma = basis_text[last_comma_pos:].strip()
                            # 如果逗号后面是新的语义段落（不是继续列举文件），则在逗号处截断
                            if after_comma and not after_comma[0] in '《【':
                                basis_text = basis_text[:last_comma_pos - 1]

                    basis_text = basis_text.strip()
                    if basis_text:
                        return basis_text

        return None

    def _extract_project_name_from_first_para(self) -> Optional[str]:
        """
        从第一段话中提取项目名称（主方法）

        句子结构：`南宁监督站****对【项目名称】****标****、****、****、****，****标*******、*******等工点`

        识别规则：
        1. 项目名称位于"对"之后
        2. 包含"铁路"两个字
        3. 在标段编号（如"HBZQ-1标"、"LWZQ-8标"）之前（允许中间有"工程"、"项目"等词）
        4. 支持"铁路广西段"这样的完整项目名称
        5. 提取结果去掉开头可能误匹配的"对"字

        示例：
        - 南宁监督站蒋德义、卢浩对【柳梧铁路】LWZQ-8标...
        - 南宁监督站李规录、陈胜对【黄百铁路广西段】HBZQ-1标...
        - ...对新建【合湛铁路】工程开展了监督检查。根据...抽查了HZZQ-1标...
        """
        # 标段编号正则：如 HBZQ-1标、LWZQ-8标、QFSG1标
        _SECTION_PAT = r'[A-Z]{2}[A-Z]+(?:-?\d+)?标'

        for para in self.paragraphs[:10]:
            if '对' not in para or '铁路' not in para:
                continue

            # 在整段文字中搜索"xxx铁路[可选后缀]"，且其后（允许间隔任意非换行字符）出现标段编号
            # 使用非贪婪匹配，取最短的铁路名称
            match = re.search(
                r'([^，。；\s对]+铁路(?:广西段|工程|项目)?)'   # 项目名称（不含"对"字）
                r'(?=[^，。；]{0,30}' + _SECTION_PAT + r')',   # lookahead：30字内出现标段编号
                para
            )
            if match:
                project_name = match.group(1).strip().lstrip('对')
                return project_name

        return None

    def _identify_section(self, para: str) -> Optional[str]:
        """
        识别段落所属的章节

        Returns:
            'rectification' - 下发整改通知单章节
            'other' - 其它问题章节
            None - 其它章节
        """
        # 下发整改通知单章节（第二章）
        if '二、' in para and ('下发整改通知单' in para or '不良行为' in para or '工点及问题' in para):
            return 'rectification'
        # 其它问题章节（第三章或第二章的"主要质量安全问题"）
        elif ('三、' in para or '其他主要' in para or '其它主要' in para or
              ('二、' in para and ('主要质量安全问题' in para or '主要安全质量问题' in para or
                                   '主要质量安全等问题' in para or '主要安全质量等问题' in para or
                                   '存在的主要' in para))) and \
             ('其他' in para or '其它' in para or '主要' in para or '问题' in para):
            return 'other'
        # 无章节编号的"存在的主要...问题"标题（如钦防二线格式）
        elif re.match(r'^存在的主要', para) and '问题' in para:
            return 'other'
        # 监督意见/有关要求章节（第三/四章）- 这些章节标志着问题部分的结束
        elif '四、' in para or '三、' in para or '监督意见' in para or '监督有关' in para or '有关要求' in para:
            return None

        return None

    def _extract_rectification_notices(self) -> List[Dict]:
        """
        提取下发整改通知单的问题

        支持两种格式：

        格式1（柳梧）：
        （一）由[施工单位]施工、[监理单位]监理的[标段名称][工点名称]（检查日期：[日期]）
        检查情况：[问题描述]
        处理措施：[处理内容]

        格式2（黄百）：
        1.中铁五局施工、西南交大监理的HBZQ-2标布柳河特大桥（检查时间：2025年9月10日）
        [问题描述]
        处理措施：[处理内容]

        返回：包含标段、工点、问题的问题列表
        """
        issues = []
        in_rectification = False
        current_section_code = None
        current_section_name = None
        current_contractor = None
        current_supervisor = None
        current_site_name = None
        current_inspection_date = None
        current_description = None
        current_requirements = None
        current_deadline = None
        current_responsible_unit = None

        # 检测文档格式
        doc_format = self._detect_document_format()

        # 提前获取检查单位和检查人员（用于所有问题）
        inspection_unit = self._extract_inspection_unit_from_first_para()
        inspection_personnel = self._extract_inspection_personnel_from_first_para()

        for i, para in enumerate(self.paragraphs):
            # 检查是否进入新章节
            section = self._identify_section(para)
            if section == 'rectification':
                in_rectification = True
                continue
            elif section == 'other':
                # 进入其它问题章节，停止收集整改通知单
                in_rectification = False
                break

            # 如果在整改通知单章节
            if in_rectification:
                # 格式1：检查是否是新标段/工点（以（一）、（二）等开头）
                if re.match(r'^（[一二三四五六七八九十]）', para):
                    # 保存前一个问题
                    if current_description:
                        issue = {
                            'section_code': current_section_code,
                            'section_name': current_section_name,
                            'site_name': current_site_name,
                            'contractor': current_contractor,
                            'supervisor': current_supervisor,
                            'inspection_unit': inspection_unit,
                            'inspection_personnel': inspection_personnel,
                            'inspection_date': current_inspection_date,
                            'description': current_description,
                            'rectification_requirements': current_requirements,
                            'rectification_deadline': current_deadline,
                            'responsible_unit': current_responsible_unit,
                            'is_rectification_notice': True,
                            'is_bad_behavior_notice': False,
                            'document_section': 'rectification'
                        }
                        issues.append(issue)

                    # 解析新的标段/工点信息
                    current_section_code = self._extract_section_code(para)
                    current_section_name = self._extract_section_name(para)
                    current_site_name = self._extract_site_name(para)
                    current_contractor = self._extract_contractor(para)
                    current_supervisor = self._extract_supervisor(para)
                    current_inspection_date = self._extract_check_date_from_para(para)
                    current_description = None
                    current_requirements = None
                    current_deadline = None
                    current_responsible_unit = None

                # 格式2：检查是否是黄百格式的数字编号行（如"1.中铁五局施工..."）
                elif doc_format == 'format2' and re.match(r'^\d+[\.．]', para):
                    # 检查是否是黄百格式的标段/工点行（包含"施工"、"监理"、"标"）
                    if '施工' in para and '监理' in para and '标' in para:
                        # 保存前一个问题
                        if current_description:
                            issue = {
                                'section_code': current_section_code,
                                'section_name': current_section_name,
                                'site_name': current_site_name,
                                'contractor': current_contractor,
                                'supervisor': current_supervisor,
                                'inspection_unit': inspection_unit,
                                'inspection_personnel': inspection_personnel,
                                'inspection_date': current_inspection_date,
                                'description': current_description,
                                'rectification_requirements': current_requirements,
                                'rectification_deadline': current_deadline,
                                'responsible_unit': current_responsible_unit,
                                'is_rectification_notice': True,
                                'is_bad_behavior_notice': '不良行为' in (current_requirements or ''),
                                'document_section': 'rectification'
                            }
                            issues.append(issue)

                        # 从数字编号行提取信息（支持单行和跨行格式）
                        contractor, supervisor, section_code, site_name, check_date = self._extract_info_from_numbered_line(para)

                        # 如果单行匹配失败，尝试跨行匹配
                        if not section_code and i > 0:
                            contractor, supervisor, section_code, site_name, check_date = self._extract_info_cross_line(self.paragraphs[i-1], para)

                        current_contractor = contractor
                        current_supervisor = supervisor
                        current_section_code = section_code
                        current_site_name = site_name
                        current_inspection_date = check_date
                        current_section_name = f"{section_code}标" if section_code else None
                        current_description = None
                        current_requirements = None
                        current_deadline = None
                        current_responsible_unit = None
                    else:
                        # 这可能是问题描述（黄百格式中，问题描述可能直接跟在工点行后面）
                        if current_site_name is not None:
                            # 去除开头的数字编号（支持半角/全角点）
                            current_description = re.sub(r'^\d+[\.．]', '', para).strip()
                        else:
                            # 如果还没有工点信息，这可能是其他内容
                            pass

                # 检查是否是"检查情况："段落
                elif para.startswith('检查情况：'):
                    current_description = para.replace('检查情况：', '').strip()

                # 检查是否是"处理措施："段落
                elif para.startswith('处理措施：'):
                    measures = para.replace('处理措施：', '').strip()
                    current_requirements = measures

                    # 添加调试日志
                    print(f"[DEBUG] 遇到'处理措施：'段落")
                    print(f"[DEBUG]   current_section_code: {current_section_code}")
                    print(f"[DEBUG]   current_description: {current_description[:50] if current_description else None}...")

                    # 从处理措施中提取整改期限
                    deadline = self._extract_deadline_from_measures(measures)
                    if deadline:
                        current_deadline = deadline

                    # 从处理措施中提取责任单位
                    responsible = self._extract_responsible_unit_from_measures(measures)
                    if responsible:
                        current_responsible_unit = responsible

                    # 判断是否是不良行为通知单
                    if '不良行为' in measures:
                        # 这是一个不良行为通知单，需要创建问题
                        if current_description:
                            print(f"[DEBUG]   → 创建问题（不良行为通知单）")
                            issue = {
                                'section_code': current_section_code,
                                'section_name': current_section_name,
                                'site_name': current_site_name,
                                'contractor': current_contractor,
                                'supervisor': current_supervisor,
                                'inspection_unit': inspection_unit,
                                'inspection_personnel': inspection_personnel,
                                'inspection_date': current_inspection_date,
                                'description': current_description,
                                'rectification_requirements': current_requirements,
                                'rectification_deadline': current_deadline,
                                'responsible_unit': current_responsible_unit,
                                'is_rectification_notice': True,
                                'is_bad_behavior_notice': True,
                                'document_section': 'rectification'
                            }
                            issues.append(issue)
                            print(f"[DEBUG]   → 问题已添加，当前总数: {len(issues)}")
                    else:

                        # 普通整改通知单：每个工点只创建一个问题，无论有多少份通知单
                        if current_description:
                            print(f"[DEBUG]   → 创建问题（普通整改通知单）")
                            issue = {
                                'section_code': current_section_code,
                                'section_name': current_section_name,
                                'site_name': current_site_name,
                                'contractor': current_contractor,
                                'supervisor': current_supervisor,
                                'inspection_unit': inspection_unit,
                                'inspection_personnel': inspection_personnel,
                                'inspection_date': current_inspection_date,
                                'description': current_description,
                                'rectification_requirements': current_requirements,
                                'rectification_deadline': current_deadline,
                                'responsible_unit': current_responsible_unit,
                                'is_rectification_notice': True,
                                'is_bad_behavior_notice': False,
                                'document_section': 'rectification'
                            }
                            issues.append(issue)
                            print(f"[DEBUG]   → 问题已添加，当前总数: {len(issues)}")

                    # 重置
                    current_description = None
                    current_requirements = None
                    current_deadline = None
                    current_responsible_unit = None

                # 黄百格式：如果当前有工点信息但没有问题描述，且这不是特殊段落，则作为问题描述
                # 注意：需要排除图片标注行（如"图1"、"图1                               图2"等）
                elif doc_format == 'format2' and current_site_name is not None and current_description is None:
                    if (not para.startswith('处理措施：') and
                        not para.startswith('检查情况：') and
                        not self._is_pic_caption(para)):
                        # 这是问题描述
                        current_description = para

        # 循环结束后，添加最后一个问题（如果还没有被添加）
        if current_description:
            issue = {
                'section_code': current_section_code,
                'section_name': current_section_name,
                'site_name': current_site_name,
                'contractor': current_contractor,
                'supervisor': current_supervisor,
                'inspection_unit': inspection_unit,
                'inspection_personnel': inspection_personnel,
                'inspection_date': current_inspection_date,
                'description': current_description,
                'rectification_requirements': current_requirements,
                'rectification_deadline': current_deadline,
                'responsible_unit': current_responsible_unit,
                'is_rectification_notice': True,
                'is_bad_behavior_notice': '不良行为' in (current_requirements or ''),
                'document_section': 'rectification'
            }
            issues.append(issue)

        return issues

    def _extract_check_date_from_para(self, para: str) -> Optional[str]:
        """
        从单个段落中提取检查日期

        格式示例：2025年5月21日
        """
        date_patterns = [
            r'(\d{4})年(\d{1,2})月(\d{1,2})日',
        ]

        for pattern in date_patterns:
            match = re.search(pattern, para)
            if match:
                year, month, day = match.groups()
                return f"{year}-{month.zfill(2)}-{day.zfill(2)}"

        return None

    def _detect_document_format(self) -> str:
        """
        检测文档格式

        返回：
        - 'format1'：有一级编号（（一）、（二）等）- 柳梧格式
        - 'format2'：无一级编号，直接使用数字编号 - 黄百格式
        """
        # 查找"二、下发整改通知单"或"三、其他"章节
        in_section = False
        for para in self.paragraphs:
            if '二、' in para or '三、' in para:
                in_section = True
                continue

            if in_section:
                # 检查是否有一级编号
                if re.match(r'^（[一二三四五六七八九十]）', para):
                    return 'format1'
                # 检查是否有数字编号（且包含标段信息）
                elif re.match(r'^\d+[\.．]', para) and ('标' in para or '施工' in para):
                    return 'format2'
                # 如果找到了其他章节标记，停止检测

                elif re.match(r'^[四五六七八九十]、', para):
                    break


        # 默认返回 format1
        return 'format1'

    def _extract_section_code(self, para: str) -> Optional[str]:
        """
        提取标段编号，如 LWZF-2, LWXQ, LWZQ-8, HBZQ-1, QFSG1 等

        支持的格式：
        - HBZQ-1 (带连字符)
        - QFSG1 (无连字符)
        - LWZQ-8 (带连字符)
        - HBZF (纯字母)

        使用前瞻断言 (?=标) 确保只匹配"标"字之前的编号，避免误匹配其他大写字母组合
        """
        # 使用前瞻断言，匹配"标"字前的标段编号
        # 模式：至少2个大写字母开头，后跟字母或数字，可选的连字符和数字
        match = re.search(r'([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?)(?=标)', para)
        if match:
            return match.group(1)
        return None

    def _extract_section_name(self, para: str) -> Optional[str]:
        """
        提取标段名称（包含标段编号和"标"字）

        支持的格式：
        - LWZF-2标
        - YCZQ-3标
        - HBZQ-1标
        - 等所有标段格式
        """
        # 使用通用模式：至少2个大写字母开头，后跟字母或数字，可选的连字符和数字，然后是"标"字
        match = re.search(r'([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?标)', para)
        if match:
            return match.group(1)
        return None

    def _extract_site_name(self, para: str) -> Optional[str]:
        """
        提取工点名称

        兼容多种时间标签格式：
        - （检查时间：2025年...）
        - （检查时间2025年...）
        - （检查日期：2025年...）
        - （检查日期2025年...）
        支持全角括号
        """
        # 统一在“标”和“检查时间/日期”之间提取工点名称，允许可选冒号
        match = re.search(r'标(.+?)（检查(?:时间|日期)[：:]?', para)
        if match:
            return match.group(1)

        return None

    def _clean_site_name_and_extract_date(self, site_name: str, current_date: Optional[str] = None) -> Tuple[str, Optional[str]]:
        """
        清理工点名称中的检查时间部分，并提取检查时间

        兼容两种格式：
        1. 检查时间在标段后面（原有格式）：工点名称不包含检查时间
        2. 检查时间在工点名称后面（新格式）：工点名称包含检查时间，需要分离

        Args:
            site_name: 原始工点名称（可能包含检查时间）
            current_date: 当前已提取的检查时间（如果有）

        Returns:
            (清理后的工点名称, 检查时间)
        """
        if not site_name:
            return site_name, current_date

        # 检查是否包含检查时间/日期
        # 支持格式：（检查时间：2026年1月21日）、（检查时间2026年1月21日）、(检查日期：...)等
        match = re.search(r'[（(]检查(?:时间|日期)[：:]?\s*(\d{4})年(\d{1,2})月(\d{1,2})日[）)]?', site_name)

        if match:
            # 提取检查时间
            year, month, day = match.groups()
            extracted_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"

            # 清理工点名称：去掉检查时间部分
            clean_name = re.sub(r'[（(]检查(?:时间|日期)[：:]?\s*\d{4}年\d{1,2}月\d{1,2}日[）)]?', '', site_name).strip()

            # 如果当前没有检查时间，使用提取的时间；否则保留原有时间
            final_date = current_date if current_date else extracted_date

            return clean_name, final_date

        # 不包含检查时间，返回原始值
        return site_name, current_date

    def _clean_site_name(self, site_name: str) -> str:
        """
        清理工点名称，去除检查时间部分

        例如：
        - "大车山中桥无砟轨道工程（检查时间2026年1月21日）" -> "大车山中桥无砟轨道工程"
        - "双贵顶隧道出口（检查时间：2026年1月21日）" -> "双贵顶隧道出口"
        """
        if not site_name:
            return site_name

        # 去除检查时间/日期部分（支持全角和半角括号，支持有无冒号）
        # 模式：（检查时间：2026年1月21日） 或 (检查时间2026年1月21日)
        cleaned = re.sub(r'[（(]检查(?:时间|日期)[：:]?[^）)]*[）)]?$', '', site_name)
        return cleaned.strip()

    def _extract_info_from_numbered_line(self, para: str) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str], Optional[str]]:
        """
        从黄百格式的数字编号行中提取信息

        格式：1.中铁五局施工、西南交大监理的HBZQ-2标布柳河特大桥（检查时间：2025年9月10日）
        或：11.中铁十四局施工、广西宁铁监理的HBZQ-4标凌云隧道2斜小里程(检查时间：2025年9月11日)。

        返回：(施工单位, 监理单位, 标段编号, 工点名称, 检查日期)
        """
        # 正则表达式：支持全角/半角编号点、顿号/逗号分隔、可选冒号、通用标段编号
        # 模式1：常规形式（监理的 后面跟 “编号标 + 工点名”）- 全角括号
        pattern = r'^\d+[\.．、]\s*(.+?)施工[、，]\s*(.+?)监理的([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?)(?=标)标(.+?)（检查(?:时间|日期)[：:]?\s*(.+?)）'
        match = re.search(pattern, para)

        # 模式2：常规形式（半角括号）
        if not match:
            pattern = r'^\d+[\.．、]\s*(.+?)施工[、，]\s*(.+?)监理的([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?)(?=标)标(.+?)\(检查(?:时间|日期)[：:]?\s*(.+?)\)'
            match = re.search(pattern, para)

        # 模式3：承包/监理单位自带标段编号，后面无“编号标”再出现（全角括号）
        # 例如：1.中铁五局YCZQ-4标施工、中铁路安YCJL-2标监理的DK262+...段路基工程（检查时间：2025年7月23日）
        if not match:
            pattern = r'^\d+[\.．]\s*(.+?)([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?)(?=标)标施工[、，]\s*(.+?)([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?)(?=标)标监理的(.+?)（检查(?:时间|日期)[：:]?\s*(.+?)）'
            match = re.search(pattern, para)
            if match:
                contractor = match.group(1).strip()
                section_code = match.group(2).strip()  # 优先使用施工单位的标段编号
                supervisor = match.group(3).strip()
                site_name = match.group(5).strip()
                check_date_str = match.group(6).strip()
                # 解析检查日期
                check_date = None
                date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', check_date_str)
                if date_match:
                    year, month, day = date_match.groups()
                    check_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                return contractor, supervisor, section_code, site_name, check_date

        # 模式4：同模式3的半角括号版本
        if not match:
            pattern = r'^\d+[\.．]\s*(.+?)([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?)(?=标)标施工[、，]\s*(.+?)([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?)(?=标)标监理的(.+?)\(检查(?:时间|日期)[：:]?\s*(.+?)\)'
            match = re.search(pattern, para)
            if match:
                contractor = match.group(1).strip()
                section_code = match.group(2).strip()
                supervisor = match.group(3).strip()
                site_name = match.group(5).strip()
                check_date_str = match.group(6).strip()
                check_date = None
                date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', check_date_str)
                if date_match:
                    year, month, day = date_match.groups()
                    check_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                return contractor, supervisor, section_code, site_name, check_date

        if match:
            contractor = match.group(1).strip()
            supervisor = match.group(2).strip()
            section_code = match.group(3).strip()
            site_name = match.group(4).strip()
            check_date_str = match.group(5).strip()

            # 解析检查日期
            check_date = None
            date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', check_date_str)
            if date_match:
                year, month, day = date_match.groups()
                check_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"

            return contractor, supervisor, section_code, site_name, check_date

        return None, None, None, None, None

    def _extract_info_cross_line(self, para1: str, para2: str) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[str], Optional[str]]:
        """
        跨行提取信息，处理标段信息在一行、工点和检查时间在下一行的情况

        格式示例：
        第一行：（一）中铁五局施工、中铁路安监理的YCZQ-4标
        第二行：1、路基DK262+635.41～DK263+079.5段（检查时间2025年7月23日）

        返回：(施工单位, 监理单位, 标段编号, 工点名称, 检查日期)
        """
        # 从第一行提取标段信息
        section_patterns = [
            # 带"由"字版本：（一）由中建八局施工、甘肃铁科监理的LWZQ-8标
            r'^（[一二三四五六七八九十]+）\s*由(.+?)施工[、，]\s*(.+?)监理的([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?)(?=标)标\s*$',
            # 无"由"字版本：（一）中铁五局施工、中铁路安监理的YCZQ-4标
            r'^（[一二三四五六七八九十]+）\s*(.+?)施工[、，]\s*(.+?)监理的([A-Z]{2,}[A-Z0-9]*(?:-?\d+)?)(?=标)标\s*$'
        ]

        contractor = None
        supervisor = None
        section_code = None

        for pattern in section_patterns:
            match = re.search(pattern, para1)
            if match:
                if len(match.groups()) == 4:  # 带"由"字版本
                    contractor = match.group(2).strip()
                    supervisor = match.group(3).strip()
                    section_code = match.group(4).strip()
                else:  # 无"由"字版本
                    contractor = match.group(1).strip()
                    supervisor = match.group(2).strip()
                    section_code = match.group(3).strip()
                break

        if not section_code:
            return None, None, None, None, None

        # 从第二行提取工点名称和检查时间
        site_patterns = [
            # 全角括号版本：1、路基DK262+635.41～DK263+079.5段（检查时间2025年7月23日）
            r'^\d+[\.．、]\s*(.+?)（检查(?:时间|日期)[：:]?\s*(.+?)）',
            # 半角括号版本
            r'^\d+[\.．、]\s*(.+?)\(检查(?:时间|日期)[：:]?\s*(.+?)\)'
        ]

        site_name = None
        check_date = None

        for pattern in site_patterns:
            match = re.search(pattern, para2)
            if match:
                site_name = match.group(1).strip()
                check_date_str = match.group(2).strip()

                # 解析检查日期
                date_match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', check_date_str)
                if date_match:
                    year, month, day = date_match.groups()
                    check_date = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                break

        if site_name and check_date:
            return contractor, supervisor, section_code, site_name, check_date

        return None, None, None, None, None

    def _extract_point_name(self, para: str) -> Optional[str]:
        """提取工点名称（兼容旧方法）"""
        return self._extract_site_name(para)

    def _extract_deadline_from_measures(self, measures: str) -> Optional[str]:
        """从处理措施中提取整改期限"""
        # 格式：2025年5月24日前完成整改
        match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', measures)
        if match:
            year, month, day = match.groups()
            return f"{year}-{month.zfill(2)}-{day.zfill(2)}"
        return None

    def _extract_responsible_unit_from_measures(self, measures: str) -> Optional[str]:
        """从处理措施中提取责任单位"""
        # 查找"施工单位"、"监理单位"等
        if '施工单位' in measures:
            return '施工单位'
        elif '监理单位' in measures:
            return '监理单位'
        elif '设计单位' in measures:
            return '设计单位'
        return None

    @staticmethod
    def _is_pic_caption(text: str) -> bool:
        """
        判断一个段落是否为纯图片标注行，应被排除在问题记录之外。

        匹配规则：段落去除空白后，仅由"图N"（N为数字）和空白字符组成。
        例如：
          - "图1"                          → True
          - "图2                           图3"  → True（多图并排，中间有大量空格）
          - "图10"                         → True
          - "存在质量隐患（图1、图2）"      → False（正常问题描述中引用图片）
          - "图1施工现场"                   → False（图片说明文字，不是纯标注）
        """
        # 去掉所有空白后，检查是否只剩"图N图M..."这样的纯图片编号序列
        collapsed = re.sub(r'\s+', '', text)
        return bool(re.fullmatch(r'(?:图[0-9]+)+', collapsed))

    def _extract_contractor(self, para: str) -> Optional[str]:
        """
        提取施工单位

        支持多种格式：
        - 格式1：由XXX施工、YYY监理的...
        - 格式2：1.XXX施工、/1．XXX施工，YYY监理的...（阿拉伯数字序号）
        - 格式3：（一）XXX施工、YYY监理的...（中文序号）
        - 格式4：施工单位：XXX
        - 格式5：XXX施工、YYY监理的...（直接以施工单位开头）
        """
        # 优先匹配“由XXX施工”形式
        match = re.search(r'由(.+?)施工', para)
        if match:
            return match.group(1).strip()

        # 兼容数字编号行开头：1.XXX施工 或 1．XXX施工
        match = re.search(r'^\d+[\.．]\s*([^、，]+?)施工', para)
        if match:
            return match.group(1).strip()

        # 新增：支持中文序号开头：（一）XXX施工、YYY监理
        # 匹配 （一）、（二）、（三）等格式，支持全角和半角括号
        match = re.search(r'^[（(][一二三四五六七八九十百]+[）)]\s*([^、，]+?)施工', para)
        if match:
            return match.group(1).strip()

        # 匹配"施工单位：XXX"格式
        match = re.search(r'施工单位[：:]\s*(.+?)(?:[，、；;]|$)', para)
        if match:
            return match.group(1).strip()

        # 匹配直接以施工单位开头的格式：XXX施工、YYY监理
        # 但要排除以数字、括号等开头的情况
        match = re.search(r'^([^0-9（）\(\)一二三四五六七八九十]+?)施工[、，]', para)
        if match:
            contractor = match.group(1).strip()
            # 排除过短的匹配（可能是误匹配）
            if len(contractor) >= 3:
                return contractor

        return None

    def _extract_supervisor(self, para: str) -> Optional[str]:
        """提取监理单位（支持顿号、逗号两种分隔符）"""
        match = re.search(r'[、，]([^、，]+?)监理', para)
        if match:
            return match.group(1)
        return None

    def _extract_total_issues_count(self) -> Optional[Dict]:
        """
        提取文档中声明的问题总数

        支持多种格式：
        格式1（柳梧）：检查发现各类安全质量问题53个（安全问题12个、质量问题21个、管理行为问题20个）
        格式2（黄百）：共计发现各类问题79个,其中质量问题27个，安全问题29个，管理行为及其他问题23个

        返回：
        {
            'total': 53,
            'safety': 12,
            'quality': 21,
            'management': 20
        }
        或 None 如果未找到
        """
        # 查找包含问题总数的段落
        for para in self.paragraphs[:20]:  # 只查看前 20 段
            if '发现' in para and '问题' in para:
                # 尝试多种格式的正则表达式
                total_count = None

                # 格式1：检查发现各类...问题(\d+)个
                match = re.search(r'检查发现各类.*?问题(\d+)个', para)
                if match:
                    total_count = int(match.group(1))

                # 格式2：共计发现各类问题(\d+)个 或 发现各类问题(\d+)个
                if not match:
                    match = re.search(r'(?:共计)?发现各类问题(\d+)个', para)
                    if match:
                        total_count = int(match.group(1))

                if total_count:
                    result = {'total': total_count}

                    # 提取安全问题数
                    safety_match = re.search(r'安全问题(\d+)个', para)
                    if safety_match:
                        result['safety'] = int(safety_match.group(1))

                    # 提取质量问题数
                    quality_match = re.search(r'质量问题(\d+)个', para)
                    if quality_match:
                        result['quality'] = int(quality_match.group(1))

                    # 提取管理行为问题数（支持多种格式）
                    # 格式1：管理行为问题(\d+)个
                    # 格式2：管理行为及其他问题(\d+)个 或 管理行为及其它问题(\d+)个
                    management_match = re.search(r'管理行为(?:及其[他它])?问题(\d+)个', para)
                    if management_match:
                        result['management'] = int(management_match.group(1))

                    return result

        return None

    def _extract_other_issues(self) -> List[Dict]:
        """
        提取其它安全质量问题

        支持多种格式：

        格式1（柳梧）：标段和工点分开
        （一）由[施工单位]施工、[监理单位]监理的[标段名称]（检查日期：[日期]）
        1. [工点名称]
        （1）[问题描述]
        （2）[问题描述]

        格式2（柳梧）：标段和工点合并（单工点情况）
        （一）由[施工单位]施工、[监理单位]监理的[标段名称][工点名称]（检查日期：[日期]）
        （1）[问题描述]
        （2）[问题描述]

        格式3（黄百）：无一级编号，直接使用数字编号
        1.中铁上海局施工、内蒙古沁原监理的HBZQ-1标幼平隧道进口（检查时间：2025年9月9日）
        （1）[问题描述]
        （2）[问题描述]

        返回：包含标段、工点、问题的问题列表
        """
        issues = []
        in_other = False
        in_problem_list = False  # 标记是否已经进入问题列表（用于三级结构）
        current_section_code = None
        current_section_name = None
        current_contractor = None
        current_supervisor = None
        current_inspection_date = None
        current_site_name = None

        # 检测文档格式
        doc_format = self._detect_document_format()

        # 提前获取检查单位和检查人员（用于所有问题）
        inspection_unit = self._extract_inspection_unit_from_first_para()
        inspection_personnel = self._extract_inspection_personnel_from_first_para()

        for idx, para in enumerate(self.paragraphs):
            # 获取对应的段落对象，用于检查Word格式属性
            para_obj = self.paragraph_objects[idx] if idx < len(self.paragraph_objects) else None

            # 检查是否进入新章节
            section = self._identify_section(para)
            if section == 'other':
                in_other = True
                continue
            elif section == 'rectification':
                # 不应该回到整改通知单章节
                continue

            # 如果进入其它章节（如"三、有关要求"或"四、监督意见"），停止收集
            # 但要排除"三、其它问题"这样的章节
            if in_other and (re.match(r'^[四五六七八九十]、', para) or
                            ('三、' in para and ('有关要求' in para or '监督意见' in para)) or
                            '监督意见' in para):
                break

            # 如果尚未显式进入“其它问题”章节，但遇到一级标段行（（一）...施工、...监理的），则进入该章节
            if not in_other and re.match(r'^（[一二三四五六七八九十]）', para) and '施工' in para and '监理' in para:
                in_other = True
                # 进入下一轮循环，在 in_other 分支中处理该段落
                continue

            # 同上：数字序号格式的标段行（1. 2. 3. 格式，包含施工和监理信息）也触发进入
            if not in_other and re.match(r'^\d+[\.．]', para) and '施工' in para and '监理' in para and '标' in para:
                in_other = True
                # 进入下一轮循环，在 in_other 分支中处理该段落
                continue

            # 如果在其它问题章节
            if in_other:
                # 检查段落是否有Word自动编号
                has_word_numbering = False
                if para_obj is not None:
                    pPr = para_obj._element.get_or_add_pPr()
                    numPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                    has_word_numbering = numPr is not None

                # 检查是否是新标段（以（一）、（二）等开头，且包含"由...施工、...监理的"）
                # 这是标段编号行，不是问题编号行
                if re.match(r'^（[一二三四五六七八九十]）', para) and '施工' in para and '监理' in para:
                    # 解析新的标段信息
                    current_section_code = self._extract_section_code(para)
                    current_section_name = self._extract_section_name(para)
                    current_contractor = self._extract_contractor(para)
                    current_supervisor = self._extract_supervisor(para)
                    current_inspection_date = self._extract_check_date_from_para(para)

                    # 尝试从一级编号中提取工点名称（格式2：标段和工点合并）
                    # 格式：由[施工单位]施工、[监理单位]监理的[标段名称][工点名称]（检查日期：[日期]）
                    current_site_name = self._extract_site_name(para)
                    # 如果没有提取到工点名称，说明是格式1（标段和工点分开），重置为None
                    if current_site_name is None:
                        current_site_name = None

                    # 重置问题列表标记
                    in_problem_list = False

                # 检查是否是数字项（以数字.、或全角点开头，如"1.工点名称"、"1、工点名称"或"1．问题描述"）
                elif re.match(r'^\d+[\.、．]', para):
                    # 尝试跨行组合：上一行是标段行，本行是工点+检查时间
                    if idx > 0:
                        contractor2, supervisor2, section_code2, site_name2, check_date2 = self._extract_info_cross_line(self.paragraphs[idx-1], para)
                        if section_code2:
                            current_contractor = contractor2
                            current_supervisor = supervisor2
                            current_section_code = section_code2
                            current_site_name = site_name2
                            current_inspection_date = check_date2
                            current_section_name = f"{section_code2}标"
                            # 跨行已解析完成，进入下一条继续
                            continue

                    # 格式3（黄百）：检查是否是黄百格式的数字编号行（包含标段和工点信息）
                    if doc_format == 'format2' and ('施工' in para and '监理' in para and '标' in para):
                        # 从数字编号行提取信息（支持单行和跨行格式）
                        contractor, supervisor, section_code, site_name, check_date = self._extract_info_from_numbered_line(para)

                        # 如果单行匹配失败，尝试跨行匹配
                        if not section_code and idx > 0:
                            contractor, supervisor, section_code, site_name, check_date = self._extract_info_cross_line(self.paragraphs[idx-1], para)

                        current_contractor = contractor
                        current_supervisor = supervisor
                        current_section_code = section_code
                        current_site_name = site_name
                        current_inspection_date = check_date
                        current_section_name = f"{section_code}标" if section_code else None
                    else:
                        # 格式1/2（柳梧）：提取内容
                        match = re.search(r'^\d+[\.、．](.+)$', para)
                        if match:
                            content = match.group(1).strip()

                            # 根据文档结构类型判断是工点名称还是问题描述
                            # 三级结构：数字编号行通常是工点编号行（1.工点名称）
                            # 二级结构：数字编号行通常是问题编号行（1.问题描述）

                            # 管理类工点名称的关键词列表
                            management_keywords = ['管理方面', '管理行为', '管理问题', '管理制度']

                            # 首先检查是否是管理类工点名称（优先级最高）
                            if content in management_keywords or content.endswith('管理方面'):
                                # 这是管理类工点名称（不需要清理检查时间）
                                current_site_name = content
                            # 检查是否是新的工点名称（包含"（检查时间"或"（检查日期"）
                            elif '（检查时间' in content or '（检查日期' in content:
                                # 这是新的工点名称，需要分离工点名称和检查时间
                                current_site_name, current_inspection_date = self._clean_site_name_and_extract_date(
                                    content, current_inspection_date
                                )
                            # 根据文档结构类型判断
                            elif self.document_structure == 'three_level':
                                # 三级结构：数字编号行可能是工点编号或问题编号
                                # 判断规则：
                                # 1. 如果内容很短（<30字）且不包含问题特征词汇，则是工点名称
                                # 2. 如果内容很长（>30字）或包含问题特征词汇，则是问题描述

                                # 问题特征词汇
                                problem_keywords = ['存在', '不符', '未', '缺', '破损', '脱焊', '松脱', '不足', '过大', '过小', '不到位', '隐患', '质量', '安全']
                                has_problem_keyword = any(keyword in content for keyword in problem_keywords)

                                if len(content) < 30 and not has_problem_keyword:
                                    # 这是工点名称，清理可能包含的检查时间
                                    current_site_name, current_inspection_date = self._clean_site_name_and_extract_date(
                                        content, current_inspection_date
                                    )
                                else:
                                    # 这是问题描述
                                    in_problem_list = True
                                    # 使用上下文中的单位信息，不从问题描述段落中提取
                                    # 避免将问题描述中的"施工"、"监理"等词误识别为单位名称
                                    final_contractor = current_contractor
                                    final_supervisor = current_supervisor

                                    issue = {
                                        'section_code': current_section_code,
                                        'section_name': current_section_name,
                                        'site_name': current_site_name,
                                        'contractor': final_contractor,
                                        'supervisor': final_supervisor,
                                        'inspection_unit': inspection_unit,
                                        'inspection_personnel': inspection_personnel,
                                        'inspection_date': current_inspection_date,
                                        'description': content,
                                        'is_rectification_notice': False,
                                        'is_bad_behavior_notice': False,
                                        'document_section': 'other'
                                    }
                                    issues.append(issue)
                            elif self.document_structure == 'two_level':
                                # 二级结构：数字编号行应该是工点名称，不是问题描述
                                # 修复：在二级结构中，数字编号行（1. 2. 3.）始终是工点名称
                                # 问题描述使用三级编号（（1）（2）（3））
                                current_site_name, current_inspection_date = self._clean_site_name_and_extract_date(
                                    content, current_inspection_date
                                )
                            elif not re.match(r'^（[0-9０-９]）', content) and not re.match(r'^[⑴-⑽]', content):
                                # 这是工点名称（格式1的情况），清理可能包含的检查时间
                                # 规则：不以问题编号开头（（1）、⑴等）
                                current_site_name, current_inspection_date = self._clean_site_name_and_extract_date(
                                    content, current_inspection_date
                                )
                            else:
                                # 这是问题描述（没有工点名称的情况）
                                # 创建问题记录
                                # 使用上下文中的单位信息，不从问题描述段落中提取
                                # 避免将问题描述中的"施工"、"监理"等词误识别为单位名称
                                final_contractor = current_contractor
                                final_supervisor = current_supervisor

                                issue = {
                                    'section_code': current_section_code,
                                    'section_name': current_section_name,
                                    'site_name': current_site_name,
                                    'contractor': final_contractor,
                                    'supervisor': final_supervisor,
                                    'inspection_unit': inspection_unit,
                                    'inspection_personnel': inspection_personnel,
                                    'inspection_date': current_inspection_date,
                                    'description': content,
                                    'is_rectification_notice': False,
                                    'is_bad_behavior_notice': False,
                                    'document_section': 'other'
                                }
                                issues.append(issue)

                # 检查是否是具体问题（以（1）、（2）、⑴、⑵等开头，或有Word自动编号）
                # 支持：（1）、(1)、（１）、(１)、（10）、(10)、⑴、⑵ 等格式，以及Word自动编号
                # 但要排除工点名称行（包含"（检查时间"或"（检查日期"）
                # 同时排除图片标注行（Word自动编号有时也会应用于图片标注段落）
                elif ((re.match(r'^[（(⑴-⑽]', para) or has_word_numbering)
                      and not ('（检查时间' in para or '（检查日期' in para)
                      and not self._is_pic_caption(para)):
                    # 提取问题编号和描述
                    # 支持：（1）、(1)、（１）、(１)、（10）、(10)、⑴、⑵ 等格式
                    # 先尝试括号格式
                    match = re.search(r'^[（(][0-9０-９]+[）)](.+)$', para)
                    if not match:
                        # 尝试带圈数字格式
                        match = re.search(r'^[⑴-⑽](.+)$', para)

                    # 如果没有匹配到文本编号，但有Word自动编号，则整行都是描述
                    if not match and has_word_numbering:
                        description = para
                    elif match:
                        description = match.group(1).strip()
                    else:
                        description = None

                    if description:
                        # 标记已经进入问题列表
                        in_problem_list = True
                        # 创建问题记录
                        # 使用上下文中的单位信息，不从问题描述段落中提取
                        # 避免将问题描述中的"施工"、"监理"等词误识别为单位名称
                        final_contractor = current_contractor
                        final_supervisor = current_supervisor

                        issue = {
                            'section_code': current_section_code,
                            'section_name': current_section_name,
                            'site_name': current_site_name,
                            'contractor': final_contractor,
                            'supervisor': final_supervisor,
                            'inspection_unit': inspection_unit,
                            'inspection_personnel': inspection_personnel,
                            'inspection_date': current_inspection_date,
                            'description': description,
                            'is_rectification_notice': False,
                            'is_bad_behavior_notice': False,
                            'document_section': 'other'
                        }
                        issues.append(issue)

                # 启发式规则：识别无编号的问题
                # 如果一行既没有文本编号，也没有Word编号，但在工点名称之后，且长度足够长，则认为它是问题描述
                # 注意：需要排除图片标注行（如"图1"、"图1                               图2"等）
                # 注意：需要排除"检查情况："和"处理措施："段落（这些属于下发整改通知单章节）
                elif (current_site_name is not None and
                      not re.match(r'^（[一二三四五六七八九十]）', para) and
                      not re.match(r'^\d+[\.、]', para) and
                      not re.match(r'^[（(⑴-⑽]', para) and
                      not ('（检查时间' in para or '（检查日期' in para) and
                      not self._is_pic_caption(para) and
                      not para.startswith('检查情况：') and
                      not para.startswith('处理措施：') and
                      len(para) > 20):
                    # 这是一个无编号的问题描述
                    # 使用上下文中的单位信息，不从问题描述段落中提取
                    # 避免将问题描述中的"施工"、"监理"等词误识别为单位名称
                    final_contractor = current_contractor
                    final_supervisor = current_supervisor

                    issue = {
                        'section_code': current_section_code,
                        'section_name': current_section_name,
                        'site_name': current_site_name,
                        'contractor': final_contractor,
                        'supervisor': final_supervisor,
                        'inspection_unit': inspection_unit,
                        'inspection_personnel': inspection_personnel,
                        'inspection_date': current_inspection_date,
                        'description': para,

                        'is_rectification_notice': False,
                        'is_bad_behavior_notice': False,
                        'document_section': 'other'
                    }
                    issues.append(issue)

        return issues


    # ===== 新增：问题列表施工/监理单位回填逻辑 =====
    def _postprocess_fill_units(self, result: Dict) -> None:
        """
        对所有问题（整改通知单与其它问题）进行施工/监理单位的回填：
        1) 基于 section_code 构建承包/监理单位映射
        2) 对缺失 contractor/supervisor 的问题按映射回填
        3) 若仍缺失，使用最近一次同标段的上下文值回填
        """
        rects = result.get('rectification_notices', []) or []
        others = result.get('other_issues', []) or []
        all_issues = rects + others

        # 1) 收集映射：section_code -> {contractor, supervisor}
        sec_map: Dict[str, Dict[str, str]] = {}
        for issue in all_issues:
            sc = issue.get('section_code')
            if not sc:
                continue
            c = issue.get('contractor')
            s = issue.get('supervisor')
            if not c and not s:
                continue
            entry = sec_map.get(sc, {})
            if c:
                entry['contractor'] = c
            if s:
                entry['supervisor'] = s
            sec_map[sc] = entry

        # 2) 回填函数（就地修改）
        def fill_list(lst: List[Dict]):
            # 最近一次上下文（按遍历顺序记录该列表内见到的单位信息）
            last_by_section: Dict[str, Dict[str, str]] = {}
            for item in lst:
                sc = item.get('section_code')
                # 更新最近上下文
                if sc:
                    known = {}
                    if item.get('contractor'):
                        known['contractor'] = item['contractor']
                    if item.get('supervisor'):
                        known['supervisor'] = item['supervisor']
                    if known:
                        last_by_section[sc] = {**last_by_section.get(sc, {}), **known}

                # 开始回填
                if not item.get('contractor') or not item.get('supervisor'):
                    # 先用全局映射（来自所有问题的汇总）
                    if sc and sc in sec_map:
                        entry = sec_map[sc]
                        if not item.get('contractor') and entry.get('contractor'):
                            item['contractor'] = entry['contractor']
                        if not item.get('supervisor') and entry.get('supervisor'):
                            item['supervisor'] = entry['supervisor']
                    # 再用最近上下文（同列表内最近一次的单位信息）
                    if sc and sc in last_by_section:
                        entry2 = last_by_section[sc]
                        if not item.get('contractor') and entry2.get('contractor'):
                            item['contractor'] = entry2['contractor']
                        if not item.get('supervisor') and entry2.get('supervisor'):
                            item['supervisor'] = entry2['supervisor']
                    # 仍缺失则设默认占位
                    if not item.get('contractor'):
                        item['contractor'] = '未知施工单位'
                    if not item.get('supervisor'):
                        item['supervisor'] = '未知监理单位'

        fill_list(rects)
        fill_list(others)

def parse_word_document(file_path: str) -> Dict:
    """
    解析 Word 文档的便捷函数

    Args:
        file_path: Word 文件路径

    Returns:
        解析结果
    """
    parser = WordDocumentParser(file_path)
    return parser.parse()

